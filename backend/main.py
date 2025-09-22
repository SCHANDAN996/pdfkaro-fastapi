import io
import json
import zipfile
from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
import pikepdf
import docx
import os
import logging
from pydantic import BaseModel
from typing import List, Dict
from PIL import Image

# Basic logging setup
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="PDFkaro.in Backend")

# CORS middleware to allow requests from any origin
origins = ["*"]
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Pydantic Models for Project Exporter
class FileData(BaseModel):
    path: str
    content: str

class ProcessRequest(BaseModel):
    files: List[FileData]
    output_format: str = "txt"
    include_paths: bool = True
    align_structure: bool = False


@app.get("/")
def read_root():
    """Root endpoint to check if the server is running."""
    return {"message": "PDFkaro.in Backend is running!"}


# --- PDF Tools Endpoints ---

@app.post("/api/v1/merge")
async def merge_pdfs(files: List[UploadFile] = File(...), pages_data: str = Form(...)):
    """Merges multiple PDF files based on page instructions."""
    logger.info(f"Received {len(files)} files for merging.")
    try:
        page_instructions = json.loads(pages_data)
        file_map = {file.filename: await file.read() for file in files}
        open_pdfs = {filename: pikepdf.Pdf.open(io.BytesIO(data)) for filename, data in file_map.items()}
        
        merged_pdf = pikepdf.Pdf.new()
        for instruction in page_instructions:
            source_filename = instruction['sourceFile']
            page_index = instruction['pageIndex']
            rotation = instruction.get('rotation', 0)
            
            if source_filename in open_pdfs:
                source_pdf = open_pdfs[source_filename]
                page = source_pdf.pages[page_index]
                if rotation != 0:
                    page.rotate(rotation, relative=True)
                merged_pdf.pages.append(page)

        for pdf in open_pdfs.values():
            pdf.close()
            
        output_buffer = io.BytesIO()
        merged_pdf.save(output_buffer)
        output_buffer.seek(0)
        
        branded_filename = "merged_by_PDFkaro.in.pdf"
        return StreamingResponse(
            output_buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={branded_filename}"}
        )
    except Exception as e:
        logger.error(f"Error during merging: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/v1/split")
async def split_pdf(file: UploadFile = File(...), pages_to_extract: str = Form(...)):
    """Splits a PDF into multiple pages or extracts selected pages with rotation."""
    try:
        page_instructions = json.loads(pages_to_extract)
        logger.info(f"Page instructions received: {page_instructions}")

        pdf_bytes = await file.read()
        source_pdf = pikepdf.Pdf.open(io.BytesIO(pdf_bytes))
        
        output_buffer = io.BytesIO()

        if page_instructions:
            new_pdf = pikepdf.Pdf.new()
            for instruction in page_instructions:
                index = instruction['pageIndex']
                rotation = instruction.get('rotation', 0)
                if 0 <= index < len(source_pdf.pages):
                    page = source_pdf.pages[index]
                    if rotation != 0:
                        page.rotate(rotation, relative=True)
                    new_pdf.pages.append(page)
            
            new_pdf.save(output_buffer)
            filename = "extracted_pages_by_PDFkaro.in.pdf"
            media_type = "application/pdf"
        
        else:
            with zipfile.ZipFile(output_buffer, 'w') as zip_file:
                for i, page in enumerate(source_pdf.pages):
                    dst = pikepdf.Pdf.new()
                    dst.pages.append(page)
                    page_buffer = io.BytesIO()
                    dst.save(page_buffer)
                    page_buffer.seek(0)
                    zip_file.writestr(f"page_{i+1}_by_PDFkaro.in.pdf", page_buffer.getvalue())
            
            filename = "split_files_by_PDFkaro.in.zip"
            media_type = "application/zip"

        output_buffer.seek(0)
        return StreamingResponse(
            output_buffer,
            media_type=media_type,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        logger.error(f"Error during splitting: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/v1/extract-single-page")
async def extract_single_page(file: UploadFile = File(...), page_number: int = Form(...)):
    """Extracts a single page from a PDF."""
    try:
        pdf_bytes = await file.read()
        source_pdf = pikepdf.Pdf.open(io.BytesIO(pdf_bytes))
        
        if 0 <= page_number < len(source_pdf.pages):
            new_pdf = pikepdf.Pdf.new()
            new_pdf.pages.append(source_pdf.pages[page_number])
            output_buffer = io.BytesIO()
            new_pdf.save(output_buffer)
            output_buffer.seek(0)
            
            filename = f"page_{page_number+1}_by_PDFkaro.in.pdf"
            return StreamingResponse(
                output_buffer,
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename={filename}"}
            )
        else:
            raise HTTPException(status_code=400, detail="Invalid page number.")
    except Exception as e:
        logger.error(f"Error extracting single page: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/compress")
async def compress_pdfs(files: List[UploadFile] = File(...), level: str = Form("recommended")):
    logger.info(f"Received {len(files)} files for compression with level: {level}")
    
    dpi_map = { "low": 300, "recommended": 150, "high": 72 }
    target_dpi = dpi_map.get(level, 150)

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for file in files:
            try:
                pdf_bytes = await file.read()
                source_pdf = pikepdf.Pdf.open(io.BytesIO(pdf_bytes))
                
                for page in source_pdf.pages:
                    for key, image_obj in page.images.items():
                        try:
                            raw_image = pikepdf.Raw(image_obj)
                            pil_image = Image.open(io.BytesIO(raw_image.data))
                            
                            pil_image.thumbnail((pil_image.width * target_dpi / pil_image.info.get('dpi', (72,72))[0], 
                                                 pil_image.height * target_dpi / pil_image.info.get('dpi', (72,72))[1]))
                            
                            img_byte_arr = io.BytesIO()
                            pil_image.save(img_byte_arr, format='JPEG', quality=85, optimize=True)
                            
                            new_image_obj = pikepdf.Image(img_byte_arr)
                            image_obj.write(new_image_obj.data, stream_decode_parameters=None)
                            
                        except Exception as img_e:
                            logger.warning(f"Could not compress an image in {file.filename}: {img_e}")
                            continue

                output_pdf_buffer = io.BytesIO()
                source_pdf.save(output_pdf_buffer, linearize=True, compress_streams=True, object_stream_mode=pikepdf.ObjectStreamMode.generate)
                output_pdf_buffer.seek(0)
                
                base_name = os.path.splitext(file.filename)[0]
                compressed_filename = f"{base_name}_compressed.pdf"
                
                zip_file.writestr(compressed_filename, output_pdf_buffer.getvalue())

            except Exception as pdf_e:
                logger.error(f"Failed to process file {file.filename}: {pdf_e}")
                continue

    zip_buffer.seek(0)
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": "attachment; filename=compressed_by_PDFkaro.in.zip"}
    )

# --- Project Exporter Tool Helper Functions & Endpoints ---
def create_tree_structure(files: List[FileData]):
    structure = {}
    for file in files:
        parts = file.path.split('/')
        current_level = structure
        for part in parts[:-1]:
            if part not in current_level:
                current_level[part] = {}
            current_level = current_level[part]
        current_level[parts[-1]] = None
    return structure

def generate_aligned_output(structure: Dict, path_prefix="", is_last=True):
    output = ""
    items = list(structure.items())
    for i, (name, content) in enumerate(items):
        is_current_last = (i == len(items) - 1)
        connector = "└── " if is_current_last else "├── "
        output += f"{path_prefix}{connector}{name}\n"
        if isinstance(content, dict):
            new_prefix = path_prefix + ("    " if is_current_last else "│   ")
            output += generate_aligned_output(content, new_prefix, is_current_last)
    return output

@app.post("/api/v1/process-structure")
async def process_structure(request: ProcessRequest):
    try:
        final_content = ""
        if request.align_structure:
            file_structure = create_tree_structure(request.files)
            final_content += "Project Structure:\n"
            final_content += generate_aligned_output(file_structure)
            final_content += "\n" + ("=" * 50) + "\n\n"
        for file_data in request.files:
            if request.include_paths:
                final_content += f"--- File: {file_data.path} ---\n\n"
            final_content += file_data.content + "\n\n"
            if request.include_paths:
                final_content += f"--- End of File: {file_data.path} ---\n\n" + ("=" * 50) + "\n\n"
        if request.output_format == "docx":
            doc = docx.Document()
            doc.add_paragraph(final_content)
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = "project_export_by_PDFkaro.in.docx"
        else:
            buffer = io.BytesIO(final_content.encode('utf-8'))
            media_type = "text/plain"
            filename = "project_export_by_PDFkaro.in.txt"
        return StreamingResponse(buffer, media_type=media_type, headers={"Content-Disposition": f"attachment; filename={filename}"})
    except Exception as e:
        logger.error(f"Error processing structure: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/export-zip-structure")
async def export_zip_structure(request: ProcessRequest):
    try:
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            if request.align_structure:
                tree_structure = create_tree_structure(request.files)
                summary_content = "Project Structure:\n"
                summary_content += generate_aligned_output(tree_structure)
                zip_file.writestr("00_project_structure.txt", summary_content)
            for file_data in request.files:
                file_content = ""
                if request.include_paths:
                    file_content += f"--- File: {file_data.path} ---\n\n"
                file_content += file_data.content
                if request.output_format == "docx":
                    doc = docx.Document()
                    doc.add_paragraph(file_content)
                    file_buffer = io.BytesIO()
                    doc.save(file_buffer)
                    file_buffer.seek(0)
                    base_name = os.path.splitext(file_data.path)[0]
                    file_name_in_zip = f"{base_name}.docx"
                    zip_file.writestr(file_name_in_zip, file_buffer.getvalue())
                else:
                    file_name_in_zip = f"{file_data.path}.txt" if not file_data.path.endswith('.txt') else file_data.path
                    zip_file.writestr(file_name_in_zip, file_content.encode('utf-8'))
        zip_buffer.seek(0)
        return StreamingResponse(
            zip_buffer,
            media_type="application/zip",
            headers={"Content-Disposition": "attachment; filename=project_export_by_PDFkaro.in.zip"}
        )
    except Exception as e:
        logger.error(f"Error creating zip archive: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))
