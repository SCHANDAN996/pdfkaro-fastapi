import io
import zipfile
import os
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
import docx
from pydantic import BaseModel
from typing import List, Dict
import logging

router = APIRouter()
logger = logging.getLogger(__name__)

class FileData(BaseModel):
    path: str
    content: str

class ProcessRequest(BaseModel):
    files: List[FileData]
    output_format: str = "txt"
    include_paths: bool = True
    align_structure: bool = False

def create_tree_structure(files: List[FileData]):
    structure = {}
    for file in files:
        parts = file.path.split('/')
        current_level = structure
        for part in parts[:-1]:
            if part not in current_level: current_level[part] = {}
            current_level = current_level[part]
        current_level[parts[-1]] = None
    return structure

def generate_aligned_output(structure: Dict, path_prefix="", is_last=True):
    output = ""
    items = list(structure.items())
    for i, (name, content) in enumerate(items):
        is_current_last = (i == len(items) - 1)
        connector = "└── " if is_current_last else "├── "
        output += f"{path_prefix}{connector}{name}\n"
        if isinstance(content, dict):
            new_prefix = path_prefix + ("    " if is_current_last else "│   ")
            output += generate_aligned_output(content, new_prefix, is_current_last)
    return output

@router.post("/single")
async def process_structure(request: ProcessRequest):
    try:
        final_content = ""
        if request.align_structure:
            file_structure = create_tree_structure(request.files)
            final_content += "Project Structure:\n" + generate_aligned_output(file_structure) + "\n" + ("=" * 50) + "\n\n"
        for file_data in request.files:
            if request.include_paths:
                final_content += f"--- File: {file_data.path} ---\n\n"
            final_content += file_data.content + "\n\n"
        if request.output_format == "docx":
            doc = docx.Document()
            doc.add_paragraph(final_content)
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = "project_export_by_PDFkaro.in.docx"
        else:
            buffer = io.BytesIO(final_content.encode('utf-8'))
            media_type = "text/plain"
            filename = "project_export_by_PDFkaro.in.txt"
        return StreamingResponse(buffer, media_type=media_type, headers={"Content-Disposition": f"attachment; filename={filename}"})
    except Exception as e:
        raise HTTPException(status_code=500, detail="Error creating single file export.")

@router.post("/zip")
async def export_zip_structure(request: ProcessRequest):
    try:
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            if request.align_structure:
                summary_content = "Project Structure:\n" + generate_aligned_output(create_tree_structure(request.files))
                zip_file.writestr("00_project_structure.txt", summary_content)
            for file_data in request.files:
                file_content = f"--- File: {file_data.path} ---\n\n" + file_data.content if request.include_paths else file_data.content
                if request.output_format == "docx":
                    doc = docx.Document()
                    doc.add_paragraph(file_content)
                    file_buffer = io.BytesIO()
                    doc.save(file_buffer)
                    file_buffer.seek(0)
                    base_name = os.path.splitext(file_data.path)[0]
                    file_name_in_zip = f"{base_name}.docx"
                    zip_file.writestr(file_name_in_zip, file_buffer.getvalue())
                else:
                    file_name_in_zip = f"{file_data.path}.txt" if not file_data.path.endswith('.txt') else file_data.path
                    zip_file.writestr(file_name_in_zip, file_content.encode('utf-8'))
        zip_buffer.seek(0)
        return StreamingResponse(zip_buffer, media_type="application/zip", headers={"Content-Disposition": "attachment; filename=project_export_by_PDFkaro.in.zip"})
    except Exception as e:
        raise HTTPException(status_code=500, detail="Error creating zip export.")
